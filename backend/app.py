"""
app.py — Spec Header Converter Backend
Receives .docx files, replaces header text, converts to PDF via LibreOffice,
returns a .zip of the PDFs. Deploy free on Render.com.
"""

import os, re, zipfile, tempfile, subprocess, shutil
from pathlib import Path
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document

app = Flask(__name__)
CORS(app)

LIBREOFFICE = shutil.which('libreoffice') or shutil.which('soffice') or 'libreoffice'

STATUS_VARIANTS = sorted([
    'Not for Construction', 'Issued for Review', 'Issued for Construction',
    'Reissued for Construction', 'Issued for Tender', 'Reissued for Tender',
    'For Review', 'For Construction', 'For Tender',
    'Addendum', 'For Information',
], key=len, reverse=True)

DATE_PATTERN = re.compile(
    r'\b(January|February|March|April|May|June|July|August|September|'
    r'October|November|December)\s+\d{4}(?:[.\s]*V\d+)?\b', re.IGNORECASE
)

def replace_in_paragraph(para, new_status, new_date):
    """
    Replace status and date in a paragraph, handling text split across multiple runs.
    First tries per-run replacement. Then checks the full concatenated paragraph text
    to catch variants that Word has split across run boundaries mid-word.
    """
    if not para.runs:
        return

    # Pass 1: per-run replacement (fast path, preserves all run formatting)
    for run in para.runs:
        for v in STATUS_VARIANTS:
            if re.search(re.escape(v), run.text, re.IGNORECASE):
                run.text = re.sub(re.escape(v), new_status, run.text, flags=re.IGNORECASE)
                break
        if DATE_PATTERN.search(run.text):
            run.text = DATE_PATTERN.sub(new_date, run.text)

    # Pass 2: full paragraph text check — catches variants split across runs
    full_text = ''.join(r.text for r in para.runs)
    needs_status = any(re.search(re.escape(v), full_text, re.IGNORECASE) for v in STATUS_VARIANTS)
    needs_date   = bool(DATE_PATTERN.search(full_text))

    if needs_status or needs_date:
        new_text = full_text
        if needs_status:
            for v in STATUS_VARIANTS:
                new_text, count = re.subn(re.escape(v), new_status, new_text, flags=re.IGNORECASE)
                if count:
                    break
        if needs_date:
            new_text = DATE_PATTERN.sub(new_date, new_text)

        if new_text != full_text:
            # Consolidate into first run, clear the rest
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ''


def replace_in_header(doc, new_status, new_date):
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if not header: continue
            containers = list(header.paragraphs)
            for tbl in header.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        containers += list(cell.paragraphs)
            for para in containers:
                replace_in_paragraph(para, new_status, new_date)


def to_pdf(docx_path, out_dir):
    r = subprocess.run(
        [LIBREOFFICE, '--headless', '--convert-to', 'pdf', '--outdir', str(out_dir), str(docx_path)],
        capture_output=True, text=True, timeout=120
    )
    if r.returncode != 0:
        raise RuntimeError(r.stderr.strip())
    p = Path(out_dir) / (Path(docx_path).stem + '.pdf')
    if not p.exists():
        raise RuntimeError('PDF not created')
    return p


@app.route('/convert', methods=['POST'])
def convert():
    files    = request.files.getlist('files')
    status   = request.form.get('status', 'Issued for Construction')
    date_str = request.form.get('date', '')
    if not files:
        return jsonify({'error': 'No files'}), 400

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        pdfs = []
        for f in files:
            if not f.filename.lower().endswith('.docx'): continue
            src = tmp / f.filename
            f.save(str(src))
            doc = Document(str(src))
            replace_in_header(doc, status, date_str)
            mod = tmp / ('_mod_' + f.filename)
            doc.save(str(mod))
            pdf = to_pdf(mod, tmp)
            original_pdf_name = Path(f.filename).stem + '.pdf'
            pdfs.append((pdf, original_pdf_name))

        if not pdfs:
            return jsonify({'error': 'No valid .docx files'}), 400

        zp = tmp / 'converted-pdfs.zip'
        with zipfile.ZipFile(zp, 'w', zipfile.ZIP_DEFLATED) as zf:
            for pdf_path, pdf_name in pdfs:
                zf.write(pdf_path, pdf_name)

        return send_file(str(zp), mimetype='application/zip',
                         as_attachment=True, download_name='converted-pdfs.zip')


@app.route('/health')
def health():
    return jsonify({'status': 'ok'})


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
